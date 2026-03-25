"""Generate LLD Part 1 Word document - Backend API, Database, LLM Service."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    return h

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for rd in rows:
        row = t.add_row().cells
        for i, v in enumerate(rd):
            row[i].text = str(v)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return t

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

# ==================== TITLE PAGE ====================
doc.add_paragraph()
title = doc.add_heading('Low-Level Design (LLD) Document', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in title.runs:
    r.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
sub = doc.add_heading('Part 1: Backend API, Database & LLM Service', level=2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
table(['Field', 'Value'], [
    ['Project', 'hunterAI GPO AdminFee Agent'],
    ['Document', 'LLD Part 1 of 3'],
    ['Date', 'March 2026'],
    ['Scope', 'FastAPI Backend, Database Layer, LLM Service'],
])
doc.add_page_break()

# ==================== TOC ====================
heading('Table of Contents', 1)
for item in [
    '1. FastAPI Application (main.py)',
    '2. Pydantic Request/Response Models',
    '3. Session Management',
    '4. API Endpoints - Detailed Specification',
    '5. Database Layer (database.py)',
    '6. Database Schema - Column Level',
    '7. LLM Service Layer (llm_service.py)',
    '8. Error Handling Strategy',
]:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ==================== 1. FASTAPI APP ====================
heading('1. FastAPI Application (main.py)', 1)

heading('1.1 Application Initialization', 2)
doc.add_paragraph('The FastAPI application is initialized with metadata and CORS middleware:')
code(
    'app = FastAPI(\n'
    '    title="hunterAI GPO AdminFee Reconciliation API",\n'
    '    version="2.0.0"\n'
    ')\n\n'
    'app.add_middleware(\n'
    '    CORSMiddleware,\n'
    '    allow_origins=["*"],\n'
    '    allow_credentials=True,\n'
    '    allow_methods=["*"],\n'
    '    allow_headers=["*"]\n'
    ')'
)

heading('1.2 Module Dependencies', 2)
table(['Import', 'Source', 'Purpose'], [
    ['FastAPI, HTTPException', 'fastapi', 'Web framework and error handling'],
    ['CORSMiddleware', 'fastapi.middleware.cors', 'Cross-origin request support'],
    ['BaseModel', 'pydantic', 'Request/response validation'],
    ['uuid', 'stdlib', 'Session ID generation'],
    ['run_sql, get_processing_status, ...', 'backend.database', 'Database operations'],
    ['normalize_delivery, generate_analysis_sql, ...', 'backend.llm_service', 'LLM operations'],
])

heading('1.3 Environment Loading', 2)
doc.add_paragraph(
    'The application loads environment variables from a .env file at the project root. '
    'This supports local development on Windows. Variables are loaded using os.environ.setdefault() '
    'to avoid overwriting existing environment variables.'
)

# ==================== 2. PYDANTIC MODELS ====================
heading('2. Pydantic Request/Response Models', 1)

heading('2.1 ChatMessage', 2)
doc.add_paragraph('Used by: POST /api/agent/start, POST /api/agent/chat')
table(['Field', 'Type', 'Required', 'Description'], [
    ['session_id', 'Optional[str]', 'No', 'UUID-based session identifier (8 chars)'],
    ['message', 'str', 'Yes', 'User input text'],
])

heading('2.2 StatusQuestion', 2)
doc.add_paragraph('Used by: POST /api/status/ask')
table(['Field', 'Type', 'Required', 'Description'], [
    ['question', 'str', 'Yes', 'Natural language question about status'],
    ['delivery', 'str', 'Yes', 'Delivery name to filter status'],
])

heading('2.3 AnalysisSetup', 2)
doc.add_paragraph('Used by: POST /api/analysis/setup, POST /api/reports/reconciliation')
table(['Field', 'Type', 'Required', 'Description'], [
    ['delivery', 'str', 'Yes', 'Delivery name to analyze'],
])

heading('2.4 AnalysisQuestion', 2)
doc.add_paragraph('Used by: POST /api/analysis/ask')
table(['Field', 'Type', 'Required', 'Description'], [
    ['question', 'str', 'Yes', 'Natural language question about contracts'],
    ['delivery', 'str', 'Yes', 'Delivery name context'],
    ['contracts', 'Optional[list[str]]', 'No', 'Contract list (auto-fetched if empty)'],
])

heading('2.5 SubmitContracts', 2)
doc.add_paragraph('Used by: Internal contract submission flow')
table(['Field', 'Type', 'Required', 'Default', 'Description'], [
    ['contracts', 'list[str]', 'Yes', '-', 'List of contract identifiers'],
    ['delivery_name', 'str', 'Yes', '-', 'Delivery batch name'],
    ['input_type', 'str', 'No', '"manual"', 'Input source: "manual" or "file"'],
])

# ==================== 3. SESSION MANAGEMENT ====================
heading('3. Session Management', 1)

heading('3.1 Session Store', 2)
doc.add_paragraph(
    'Sessions are stored in an in-memory Python dictionary (sessions = {}). '
    'Each session is identified by a truncated UUID (8 characters).'
)

heading('3.2 Session State Schema', 2)
table(['Field', 'Type', 'Values', 'Description'], [
    ['step', 'str', 'input_type | contracts_manual | contracts_file | delivery | confirm | trigger | done', 'Current workflow step'],
    ['input_type', 'Optional[str]', '"manual" | "file" | None', 'How contracts are provided'],
    ['contracts', 'Optional[str]', 'Raw input string', 'Original contract input'],
    ['contracts_list', 'list[str]', '[]', 'Parsed contract names'],
    ['delivery_name', 'Optional[str]', 'Any string', 'Delivery identifier'],
    ['job_status', 'Optional[str]', '"TRIGGERED" | "FAILED" | None', 'Pipeline execution status'],
])

heading('3.3 Step State Machine', 2)
doc.add_paragraph('The processing agent follows a strict linear state machine:')
code(
    'input_type ──► contracts_manual ──► delivery ──► confirm ──► trigger ──► done\n'
    '           └─► contracts_file   ──►'
)

table(['Current Step', 'Valid Inputs', 'Next Step', 'Action'], [
    ['input_type', '"manual"', 'contracts_manual', 'Set input_type'],
    ['input_type', '"file"', 'contracts_file', 'Set input_type'],
    ['contracts_manual', 'Comma-separated names', 'delivery', 'Parse and store contracts'],
    ['contracts_file', '"done"', 'delivery', 'Mark as S3 file input'],
    ['contracts_file', 'Comma-separated names', 'delivery', 'Override with manual input'],
    ['delivery', 'Any non-empty string', 'confirm', 'Store delivery name'],
    ['confirm', '"yes" / "y"', 'trigger', 'Insert into DB'],
    ['confirm', '"no" / "n" / "restart"', 'input_type', 'Reset session'],
    ['trigger', '"trigger"', 'done', 'Execute Airflow DAG'],
    ['trigger', '"skip"', 'done', 'Skip pipeline'],
    ['done', 'Any', 'done', 'Return completion message'],
])

# ==================== 4. API ENDPOINTS ====================
heading('4. API Endpoints - Detailed Specification', 1)

# --- Panel 1 ---
heading('4.1 Panel 1: Processing Agent Chat', 2)

heading('POST /api/agent/start', 3)
doc.add_paragraph('Initializes a new processing session and returns a greeting.')
table(['Property', 'Value'], [
    ['Method', 'POST'],
    ['Auth', 'None'],
    ['Request Body', 'None'],
])
doc.add_paragraph('Response (200):')
table(['Field', 'Type', 'Description'], [
    ['session_id', 'str', '8-character UUID'],
    ['agent_message', 'str', 'Greeting message with markdown'],
    ['step', 'str', '"input_type"'],
    ['expecting', 'str', '"manual or file"'],
])

heading('POST /api/agent/chat', 3)
doc.add_paragraph('Continues the conversational workflow. Response varies by current step.')
table(['Property', 'Value'], [
    ['Method', 'POST'],
    ['Auth', 'None'],
    ['Request Body', 'ChatMessage {session_id, message}'],
    ['Error 400', 'Invalid or missing session_id'],
])
doc.add_paragraph('Response (200) — fields present in all steps:')
table(['Field', 'Type', 'Description'], [
    ['session_id', 'str', 'Echo back session ID'],
    ['agent_message', 'str', 'Agent response (markdown)'],
    ['step', 'str', 'Current step name'],
    ['expecting', 'Optional[str]', 'Hint for expected input'],
])
doc.add_paragraph('Additional fields when step = "done":')
table(['Field', 'Type', 'Description'], [
    ['delivery', 'str', 'Delivery name submitted'],
    ['contracts', 'list[str]', 'Contract list submitted'],
])

# --- Panel 2 ---
heading('4.2 Panel 2: Status Monitor', 2)

heading('GET /api/status/summary', 3)
doc.add_paragraph('Returns aggregate processing status counts.')
doc.add_paragraph('Response (200):')
table(['Field', 'Type', 'Example', 'Description'], [
    ['total', 'int', '5', 'Total contracts in metadata'],
    ['completed', 'int', '3', 'Contracts with STATUS = 0'],
    ['in_progress', 'int', '2', 'Contracts with STATUS != 0'],
])

heading('GET /api/status/contracts', 3)
doc.add_paragraph('Returns contract-level detail rows.')
doc.add_paragraph('Response (200):')
table(['Field', 'Type', 'Description'], [
    ['columns', 'list[str]', '["CONTRACT_NAME", "DELIVERY", "STATUS"]'],
    ['rows', 'list[list]', 'Array of row arrays'],
])

heading('POST /api/status/ask', 3)
doc.add_paragraph('AI-powered natural language Q&A about processing status.')
doc.add_paragraph('Request: StatusQuestion {question, delivery}')
doc.add_paragraph('Response (200):')
table(['Field', 'Type', 'Description'], [
    ['question', 'str', 'Echo user question'],
    ['sql_generated', 'str', 'LLM-generated SQL query'],
    ['raw_data', 'dict', '{columns, rows} from DB'],
    ['answer', 'str', 'Human-friendly interpretation'],
])

# --- Panel 3 ---
heading('4.3 Panel 3: Contract Analysis', 2)

heading('GET /api/analysis/deliveries', 3)
doc.add_paragraph('Lists all available deliveries from audit data.')
doc.add_paragraph('Response (200):')
table(['Field', 'Type', 'Description'], [
    ['deliveries', 'list[str]', 'Sorted unique delivery names'],
])

heading('POST /api/analysis/setup', 3)
doc.add_paragraph('Normalizes delivery name via LLM and fetches associated contracts.')
doc.add_paragraph('Request: AnalysisSetup {delivery}')
doc.add_paragraph('Response (200):')
table(['Field', 'Type', 'Description'], [
    ['delivery_original', 'str', 'User-provided delivery name'],
    ['delivery_normalized', 'str', 'LLM-normalized name (e.g., delivery_10)'],
    ['contracts', 'list[str]', 'Contract names for this delivery'],
    ['message', 'str', 'Status message'],
])

heading('POST /api/analysis/ask', 3)
doc.add_paragraph('4-stage NL question answering pipeline.')
doc.add_paragraph('Request: AnalysisQuestion {question, delivery, contracts?}')
doc.add_paragraph('Processing Pipeline:')
table(['Stage', 'Action', 'LLM Call'], [
    ['1', 'Generate SQL from question + contracts', 'Yes (sql_generator.txt)'],
    ['2', 'Execute SQL against MySQL', 'No'],
    ['3', 'Analyze raw results', 'Yes (contract_analyst_prompt.txt)'],
    ['4', 'Format for business users', 'Yes (response_formatter.txt)'],
])
doc.add_paragraph('Response (200):')
table(['Field', 'Type', 'Description'], [
    ['question', 'str', 'Echo user question'],
    ['contracts', 'list[str]', 'Contracts used'],
    ['sql_generated', 'str', 'Generated MySQL query'],
    ['raw_data', 'dict', '{columns, rows}'],
    ['analysis', 'str', 'Technical analysis'],
    ['formatted_answer', 'str', 'Business-friendly answer'],
])

# --- Reports ---
heading('4.4 Reports & Reconciliation', 2)

heading('POST /api/reports/reconciliation', 3)
doc.add_paragraph('Generates PO Master vs AdminFee Report spend comparison.')
doc.add_paragraph('Request: AnalysisSetup {delivery}')
doc.add_paragraph('Response (200):')
table(['Field', 'Type', 'Description'], [
    ['delivery', 'str', 'Delivery name'],
    ['contracts', 'list[str]', 'Contracts analyzed'],
    ['summary', 'list[dict]', 'Per-contract reconciliation data'],
    ['formatted_report', 'str', 'LLM-formatted summary'],
])

doc.add_paragraph('Each item in summary list:')
table(['Field', 'Type', 'Description'], [
    ['contract', 'str', 'Contract ID'],
    ['report_spend', 'float|None', 'Sales volume from admin_fee_report'],
    ['po_spend', 'float', 'SUM(PO_Base_Spend_actual)'],
    ['inv_spend', 'float', 'SUM(INV_Extended_Spend_actual)'],
    ['selected_spend', 'float', 'MAX(po_spend, inv_spend)'],
    ['difference', 'float|None', 'ABS(selected_spend - report_spend)'],
    ['status', 'str', '"MATCH" | "PO Master is high" | "Report higher" | "Contract not found..."'],
])

heading('GET /api/reports/audit', 3)
doc.add_paragraph('Returns audit trail data. Optional query param: ?delivery=delivery_10')
doc.add_paragraph('Response (200):')
table(['Field', 'Type', 'Description'], [
    ['columns', 'list[str]', '["contract_name","delivery","excel_file_path","status","record_cnt","insrt_date"]'],
    ['rows', 'list[list]', 'Audit records (max 100 if no filter)'],
])

heading('GET /api/health', 3)
doc.add_paragraph('Response (200):')
table(['Field', 'Type', 'Value'], [
    ['status', 'str', '"healthy"'],
    ['service', 'str', '"hunterAI GPO AdminFee Reconciliation API"'],
    ['version', 'str', '"2.0"'],
])

# ==================== 5. DATABASE LAYER ====================
heading('5. Database Layer (database.py)', 1)

heading('5.1 Connection Configuration', 2)
table(['Parameter', 'Value', 'Description'], [
    ['Engine', 'SQLAlchemy create_engine', 'Connection pool manager'],
    ['Driver', 'mysql+pymysql', 'MySQL dialect with PyMySQL driver'],
    ['Host', 'dev-db-test.c969yoyq9cyy.us-east-1.rds.amazonaws.com', 'AWS RDS endpoint'],
    ['Port', '3306', 'MySQL default port'],
    ['Default DB', 'joblog_metadata', 'Primary metadata database'],
    ['pool_pre_ping', 'True', 'Validates connections before use'],
    ['pool_recycle', '3600', 'Recycle connections after 1 hour'],
])

heading('5.2 Function Signatures', 2)
table(['Function', 'Parameters', 'Returns', 'Description'], [
    ['run_sql()', 'query: str', '{columns: list, rows: list[list]}', 'Execute raw SQL query'],
    ['run_sql_params()', 'query: str, params: dict', '{columns: list, rows: list[list]}', 'Execute parameterized query'],
    ['get_contract_for_delivery()', 'delivery: str', 'list[str]', 'Get contract names for delivery'],
    ['get_processing_status()', '(none)', '{total, completed, in_progress}', 'Aggregate status counts'],
    ['get_contract_details()', '(none)', '{columns, rows}', 'All contract records'],
    ['get_deliveries()', '(none)', 'list[str]', 'Unique delivery names'],
    ['insert_contracts_metadata()', 'contracts: list, delivery: str', 'bool', 'Truncate + insert contracts'],
    ['get_contract_summary()', 'contracts: list', 'list[dict]', 'PO vs Report reconciliation'],
    ['get_audit_data()', 'delivery: Optional[str]', '{columns, rows}', 'Audit trail records'],
])

heading('5.3 Reconciliation Logic (get_contract_summary)', 2)
doc.add_paragraph('For each contract, the function:')
bullet('Converts hyphens to underscores: PP-NS-123 becomes PP_NS_123')
bullet('Queries admin_fee.PO_Master_{contract}_v1 for SUM(PO_Base_Spend_actual) and SUM(INV_Extended_Spend_actual)')
bullet('Queries admin_fee.admin_fee_report for SUM(Sales_Volume) WHERE Contract ID = contract')
bullet('Calculates selected_spend = MAX(po_spend, inv_spend)')
bullet('Compares selected_spend vs report_spend to determine status')

doc.add_paragraph('Status determination:')
table(['Condition', 'Status Value'], [
    ['report_spend is None', '"Contract not found in adminFee report"'],
    ['selected_spend > report_spend', '"PO Master is high"'],
    ['selected_spend < report_spend', '"Report higher"'],
    ['selected_spend == report_spend', '"MATCH"'],
])

# ==================== 6. DATABASE SCHEMA ====================
heading('6. Database Schema - Column Level', 1)

heading('6.1 joblog_metadata.admin_fee_metadata', 2)
table(['Column', 'Type', 'Nullable', 'Default', 'Description'], [
    ['CONTRACT_NAME', 'VARCHAR(255)', 'No', '-', 'Contract identifier (e.g., PP-NS-123)'],
    ['DELIVERY', 'VARCHAR(255)', 'No', '-', 'Delivery batch name (e.g., delivery_10)'],
    ['STATUS', 'INT', 'No', '1', 'Processing status: 1=Ready, 2=SQL Done, 0=Completed'],
])

heading('6.2 joblog_metadata.adminfee_audit_data', 2)
table(['Column', 'Type', 'Nullable', 'Description'], [
    ['contract_name', 'VARCHAR(255)', 'No', 'Contract identifier'],
    ['delivery', 'VARCHAR(255)', 'No', 'Delivery name'],
    ['excel_file_path', 'VARCHAR(500)', 'Yes', 'S3 path to generated Excel'],
    ['status', 'VARCHAR(50)', 'Yes', 'Processing status text'],
    ['record_cnt', 'INT', 'Yes', 'Number of records processed'],
    ['insrt_date', 'DATETIME', 'No', 'Record insertion timestamp'],
])

heading('6.3 admin_fee.PO_Master_{CONTRACT}_{VERSION}', 2)
doc.add_paragraph('Dynamic table — one per contract. Naming: PP-NS-123 becomes PO_Master_PP_NS_123_v1')
table(['Column', 'Type', 'Description'], [
    ['PO_Base_Spend_actual', 'DECIMAL', 'Purchase order base spend amount'],
    ['INV_Extended_Spend_actual', 'DECIMAL', 'Invoice extended spend amount'],
    ['Contracted_Supplier', 'VARCHAR', 'Supplier name'],
    ['PO_Description', 'VARCHAR', 'Purchase order description'],
    ['(additional columns)', '...', 'Various PO data fields populated by Airflow pipeline'],
])

heading('6.4 admin_fee.admin_fee_report', 2)
table(['Column', 'Type', 'Description'], [
    ['Contract ID', 'VARCHAR(255)', 'Contract identifier for join'],
    ['Sales_Volume', 'DECIMAL', 'Reported sales volume from GPO'],
    ['(additional columns)', '...', 'Various report fields'],
])

heading('6.5 joblog_metadata.metadata_table_database', 2)
table(['Column', 'Type', 'Description'], [
    ['key', 'VARCHAR', 'Configuration key name'],
    ['value', 'VARCHAR', 'Configuration value (e.g., S3_AccessKey, S3_Secret_Access_Key)'],
])

# ==================== 7. LLM SERVICE ====================
heading('7. LLM Service Layer (llm_service.py)', 1)

heading('7.1 LLM Configuration', 2)
table(['Parameter', 'Value'], [
    ['Provider', 'OpenAI'],
    ['Model', 'gpt-4.1-mini'],
    ['Integration', 'LangChain ChatOpenAI'],
    ['Auth', 'OPENAI_API_KEY environment variable'],
])

heading('7.2 Prompt Templates', 2)
table(['Prompt File', 'Variable', 'Used By', 'Purpose'], [
    ['delivery_normalizer.txt', 'DELIVERY_PROMPT', 'normalize_delivery()', 'Convert user input to delivery_N format'],
    ['sql_generator.txt', 'SQL_PROMPT', 'generate_analysis_sql()', 'NL question to MySQL SELECT'],
    ['contract_analyst_prompt.txt', 'ANALYST_PROMPT', 'analyze_results()', 'Analyze SQL result data'],
    ['response_formatter.txt', 'FORMAT_PROMPT', 'format_response()', 'Format analysis for business users'],
    ['status_tracker.txt', 'STATUS_SYSTEM_MESSAGE', 'generate_status_query(), interpret_status_result()', 'Status Q&A system prompt'],
    ['contract_summary.txt', 'SUMMARY_MESSAGE', 'format_summary()', 'Reconciliation summary formatting'],
])

heading('7.3 Function Details', 2)

heading('extract_sql_query(llm_response: str) -> str', 3)
doc.add_paragraph('Extracts SQL from LLM text output using a multi-strategy approach:')
bullet('Strategy 1: Look for ```sql ... ``` code blocks (regex)')
bullet('Strategy 2: Look for prefixes like "sql_query:", "sql:", "query:"')
bullet('Strategy 3: Find first SQL keyword (SELECT, INSERT, UPDATE, DELETE, WITH)')
bullet('Strategy 4: Extract from keyword to next semicolon')
bullet('Post-processing: Replace newlines with spaces, remove backticks')
bullet('Raises ValueError if no SQL found')

heading('normalize_delivery(user_input: str) -> str', 3)
doc.add_paragraph('Two-tier normalization:')
bullet('Primary: LLM call using DELIVERY_PROMPT template')
bullet('Fallback: Regex matching pattern delivery[_\\s]*(\\d+) → delivery_{N}')

heading('generate_analysis_sql(question, contracts) -> str', 3)
doc.add_paragraph('Formats SQL_PROMPT with {question, contracts}, invokes LLM, extracts SQL via extract_sql_query().')

heading('analyze_results(question, results) -> str', 3)
doc.add_paragraph('Formats ANALYST_PROMPT with {question, result}, invokes LLM, returns raw analysis text.')

heading('format_response(analysis) -> str', 3)
doc.add_paragraph('Formats FORMAT_PROMPT with {analysis}, invokes LLM, returns business-formatted text.')

heading('generate_status_query(question, delivery) -> str', 3)
doc.add_paragraph('Combines STATUS_SYSTEM_MESSAGE with user question and delivery, invokes LLM, extracts SQL.')

heading('interpret_status_result(question, db_result) -> str', 3)
doc.add_paragraph('Combines STATUS_SYSTEM_MESSAGE with question and SQL result, invokes LLM for human-friendly answer.')

heading('format_summary(summary_data) -> str', 3)
doc.add_paragraph('Formats SUMMARY_MESSAGE with {summary_data}, invokes LLM for reconciliation report.')

# ==================== 8. ERROR HANDLING ====================
heading('8. Error Handling Strategy', 1)

heading('8.1 API Layer (main.py)', 2)
table(['Scenario', 'HTTP Code', 'Handling'], [
    ['Invalid session_id', '400', 'HTTPException with "Invalid session" message'],
    ['No contracts found', '404', 'HTTPException with "No contracts found"'],
    ['Database error', '500', 'HTTPException with error string'],
    ['LLM error', '500', 'HTTPException with error string'],
    ['Pipeline trigger failure', '200', 'Returned in agent_message with FAILED status'],
])

heading('8.2 Database Layer (database.py)', 2)
table(['Scenario', 'Handling'], [
    ['Connection failure', 'pool_pre_ping auto-reconnects stale connections'],
    ['Query error in get_contract_summary', 'Try/except per contract, error appended to summary list'],
    ['Missing table', 'SQL error propagated to API layer'],
])

heading('8.3 LLM Layer (llm_service.py)', 2)
table(['Scenario', 'Handling'], [
    ['LLM call failure in normalize_delivery', 'Falls back to regex-based _fallback_normalize_delivery()'],
    ['No SQL in LLM response', 'ValueError raised by extract_sql_query()'],
    ['Empty LLM response', 'ValueError raised by extract_sql_query()'],
    ['API key missing', 'LangChain raises authentication error'],
])

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('End of LLD Part 1 — Backend API, Database & LLM Service')
run.bold = True
run.italic = True

# ==================== SAVE ====================
output_dir = '/home/user/adminFeeAgent2/docs'
os.makedirs(output_dir, exist_ok=True)
doc.save(os.path.join(output_dir, 'LLD_Part1_Backend.docx'))
print('LLD Part 1 saved to docs/LLD_Part1_Backend.docx')
