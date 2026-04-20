"""Convert docs/LLD_AdminFeeAgent.md to a formatted Word document."""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path(__file__).parent / "docs" / "LLD_AdminFeeAgent.md"
DST = Path(__file__).parent / "docs" / "LLD_AdminFeeAgent.docx"


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_inline_runs(paragraph, text):
    """Parse inline **bold**, *italic*, `code` into runs."""
    pattern = r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)"
    for part in re.split(pattern, text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        else:
            paragraph.add_run(part)


def parse_table(block_lines):
    """Return list of rows (list of cells) from a markdown table."""
    rows = []
    for line in block_lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[\s\-:|]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    rows = [r + [""] * (ncols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_runs(p, cell_text)
            if i == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, "1F4E79")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def convert():
    md = SRC.read_text(encoding="utf-8")
    lines = md.split("\n")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Low-Level Design (LLD)")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("AdminFee Reconciliation Agent")
    r.bold = True
    r.font.size = Pt(20)

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ver.add_run("Version 1.0  ·  2026-04-20  ·  Voya Financial")
    r.italic = True
    r.font.size = Pt(12)

    doc.add_page_break()

    i = 0
    in_code = False
    code_lines = []
    table_buffer = []

    def flush_table():
        if table_buffer:
            add_table(doc, parse_table(table_buffer))
            table_buffer.clear()

    def flush_code():
        if code_lines:
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code_lines))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F2F2F2")
            pPr.append(shd)
            code_lines.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Code fences
        if stripped.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if stripped.startswith("|"):
            table_buffer.append(stripped)
            i += 1
            continue
        else:
            flush_table()

        # Horizontal rule
        if stripped == "---":
            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            h = doc.add_heading(level=0)
            r = h.add_run(stripped[2:].lstrip("# "))
            r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif stripped.startswith("## "):
            h = doc.add_heading(level=1)
            r = h.add_run(stripped[3:])
            r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif stripped.startswith("### "):
            h = doc.add_heading(level=2)
            r = h.add_run(stripped[4:])
            r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif stripped.startswith("#### "):
            h = doc.add_heading(level=3)
            r = h.add_run(stripped[5:])
        # Blockquote
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(stripped[2:])
            r.italic = True
            r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        # Bulleted list
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:])
        # Numbered list
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s", "", stripped))
        elif stripped == "":
            pass
        else:
            p = doc.add_paragraph()
            add_inline_runs(p, stripped)

        i += 1

    flush_table()
    flush_code()

    doc.save(DST)
    print(f"Written: {DST}")


if __name__ == "__main__":
    convert()
