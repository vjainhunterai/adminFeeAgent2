"""Generate HLD Word document for hunterAI AdminFee Agent."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    return h

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            for p in row[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)

# ==================== TITLE PAGE ====================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('High-Level Design (HLD) Document', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

subtitle = doc.add_heading('hunterAI GPO AdminFee Reconciliation Agent', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_table(
    ['Field', 'Value'],
    [
        ['Project Name', 'hunterAI GPO AdminFee Agent'],
        ['Version', '1.0.0'],
        ['Date', 'March 2026'],
        ['Author', 'hunterAI Engineering Team'],
        ['Status', 'Development'],
    ]
)
doc.add_page_break()

# ==================== TABLE OF CONTENTS ====================
add_heading_styled('Table of Contents', 1)
toc_items = [
    '1. Introduction',
    '2. System Overview',
    '3. Architecture Diagram',
    '4. Technology Stack',
    '5. Component Overview',
    '6. Data Flow',
    '7. Database Design (High Level)',
    '8. API Design (High Level)',
    '9. Security Architecture',
    '10. Deployment Architecture',
    '11. Non-Functional Requirements',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ==================== 1. INTRODUCTION ====================
add_heading_styled('1. Introduction', 1)

add_heading_styled('1.1 Purpose', 2)
doc.add_paragraph(
    'This document describes the High-Level Design (HLD) of the hunterAI GPO AdminFee '
    'Reconciliation Agent — an AI-powered web application that automates Group Purchasing '
    'Organization (GPO) contract processing, real-time monitoring, and data analysis using '
    'conversational AI agents.'
)

add_heading_styled('1.2 Scope', 2)
scope_items = [
    'Automated contract ingestion (manual entry or S3 file upload)',
    'ETL pipeline orchestration via Apache Airflow',
    'Real-time processing status monitoring with AI-powered Q&A',
    'Contract spend reconciliation and natural language analytics',
    'Multi-panel web interface for end-to-end workflow',
]
for item in scope_items:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('1.3 Intended Audience', 2)
audiences = ['Solution Architects', 'Development Team', 'QA Engineers', 'Project Managers', 'Business Analysts']
for a in audiences:
    doc.add_paragraph(a, style='List Bullet')

add_heading_styled('1.4 Glossary', 2)
add_table(
    ['Term', 'Description'],
    [
        ['GPO', 'Group Purchasing Organization'],
        ['AdminFee', 'Administrative fee charged on GPO contracts'],
        ['PO', 'Purchase Order'],
        ['Delivery', 'A batch processing run identifier'],
        ['LangGraph', 'Agent orchestration framework by LangChain'],
        ['DAG', 'Directed Acyclic Graph (Airflow pipeline unit)'],
        ['Reconciliation', 'Comparing PO spend vs AdminFee report spend'],
    ]
)

# ==================== 2. SYSTEM OVERVIEW ====================
add_heading_styled('2. System Overview', 1)
doc.add_paragraph(
    'The system is a three-panel web application that provides an end-to-end workflow '
    'for GPO AdminFee contract reconciliation:'
)

add_table(
    ['Panel', 'Name', 'Purpose'],
    [
        ['Panel 1', 'Processing Agent', 'Step-by-step chat to submit contracts and trigger Airflow pipeline'],
        ['Panel 2', 'Status Monitor', 'Real-time progress dashboard with AI-powered Q&A'],
        ['Panel 3', 'Analysis Chat', 'Natural language contract analysis and spend reconciliation'],
    ]
)

add_heading_styled('2.1 Key Capabilities', 2)
capabilities = [
    'Conversational Contract Processing — Step-by-step guided workflow via chat interface',
    'Automated Pipeline Orchestration — Triggers Airflow DAGs for ETL processing',
    'Real-Time Status Monitoring — Auto-refreshing dashboard with AI-powered Q&A',
    'LLM-Powered Analytics — Natural language to SQL generation for contract analysis',
    'Spend Reconciliation — Automated PO vs AdminFee report comparison',
]
for c in capabilities:
    doc.add_paragraph(c, style='List Bullet')

# ==================== 3. ARCHITECTURE ====================
add_heading_styled('3. Architecture Diagram', 1)

add_heading_styled('3.1 System Architecture', 2)
doc.add_paragraph(
    'The system follows a layered architecture with the following major components:'
)

add_code_block(
    '                    End User (Browser)\n'
    '                          |\n'
    '                          | HTTP (Port 3000)\n'
    '                          v\n'
    '                 React Frontend (Vite)\n'
    '                 - AgentChatPanel\n'
    '                 - StatusMonitorPanel\n'
    '                 - AnalysisPanel\n'
    '                          |\n'
    '                          | /api/* (Proxy to :8000)\n'
    '                          v\n'
    '                 FastAPI Backend (Uvicorn)\n'
    '                 - REST Endpoints\n'
    '                 - Session Management\n'
    '                 - LLM Service Layer\n'
    '                    /        |        \\\n'
    '                   v         v         v\n'
    '           OpenAI GPT   MySQL RDS   Airflow (EC2)\n'
    '           (LangChain)  (SQLAlchemy)  (SSH/subprocess)\n'
    '                                       |\n'
    '                                    AWS S3\n'
)

add_heading_styled('3.2 Communication Patterns', 2)
add_table(
    ['From', 'To', 'Protocol', 'Purpose'],
    [
        ['Browser', 'React App', 'HTTP/3000', 'Serve frontend'],
        ['React App', 'FastAPI', 'HTTP/8000', 'REST API calls (proxied)'],
        ['FastAPI', 'OpenAI API', 'HTTPS', 'LLM inference'],
        ['FastAPI', 'MySQL RDS', 'TCP/3306', 'Database queries'],
        ['FastAPI', 'Airflow (EC2)', 'SSH/22', 'Trigger DAG execution'],
        ['FastAPI', 'AWS S3', 'HTTPS', 'File upload/download'],
    ]
)

# ==================== 4. TECH STACK ====================
add_heading_styled('4. Technology Stack', 1)

add_heading_styled('4.1 Frontend', 2)
add_table(
    ['Component', 'Technology', 'Version'],
    [
        ['UI Library', 'React', '18.3'],
        ['Build Tool', 'Vite', '5.4'],
        ['Styling', 'Custom CSS (Dark Theme)', '-'],
        ['State Management', 'React Hooks', '-'],
        ['HTTP Client', 'Fetch API', 'Native'],
    ]
)

add_heading_styled('4.2 Backend', 2)
add_table(
    ['Component', 'Technology', 'Version'],
    [
        ['Web Framework', 'FastAPI', '0.115.0'],
        ['ASGI Server', 'Uvicorn', '0.30.0'],
        ['ORM', 'SQLAlchemy', '2.0.35'],
        ['DB Driver', 'PyMySQL', '1.1.1'],
        ['LLM Integration', 'LangChain-OpenAI', '0.2.0'],
        ['Agent Framework', 'LangGraph', '0.2.0'],
        ['Data Processing', 'Pandas', '2.2.2'],
        ['Excel Handling', 'Openpyxl', '3.1.5'],
        ['Validation', 'Pydantic', '2.9.0'],
        ['Encryption', 'Cryptography (Fernet)', '43.0.0'],
        ['SSH Client', 'Paramiko', '3.4.0'],
        ['AWS SDK', 'Boto3', '1.35.0'],
        ['WebSockets', 'websockets', '12.0'],
    ]
)

add_heading_styled('4.3 Infrastructure', 2)
add_table(
    ['Component', 'Technology'],
    [
        ['Database', 'AWS RDS (MySQL)'],
        ['File Storage', 'AWS S3'],
        ['Pipeline Engine', 'Apache Airflow'],
        ['Compute', 'AWS EC2'],
        ['LLM Provider', 'OpenAI GPT-4.1-mini'],
    ]
)

# ==================== 5. COMPONENT OVERVIEW ====================
add_heading_styled('5. Component Overview', 1)

add_heading_styled('5.1 Frontend Components', 2)
add_table(
    ['Component', 'Responsibility'],
    [
        ['AgentChatPanel.jsx', 'Step-by-step chat to submit contracts and trigger pipeline'],
        ['StatusMonitorPanel.jsx', 'Auto-refreshing status cards, progress bar, contract table, AI Q&A'],
        ['AnalysisPanel.jsx', 'Delivery selection, reconciliation reports, NL question answering'],
        ['MarkdownRenderer.jsx', 'Parse markdown (bold, code, tables, lists) into HTML'],
        ['api.js', 'Centralized API call functions with error handling'],
    ]
)

add_heading_styled('5.2 Backend Modules', 2)
add_table(
    ['Module', 'Responsibility'],
    [
        ['main.py', 'FastAPI app, REST endpoints, session management'],
        ['database.py', 'SQLAlchemy connection pool, query execution, DB helpers'],
        ['llm_service.py', 'LLM orchestration — SQL generation, analysis, formatting'],
        ['prompts/', 'System prompt templates for each LLM task'],
    ]
)

add_heading_styled('5.3 CLI Agents (Standalone)', 2)
add_table(
    ['Module', 'Responsibility'],
    [
        ['adminFee_Master_agent.py', 'CLI menu — routes to processing or analysis workflow'],
        ['adminfee_processing_agent.py', 'LangGraph workflow — full processing pipeline (11 nodes)'],
        ['contract_analyst_agent_cot.py', 'LangGraph workflow — NL analysis loop (4 nodes)'],
    ]
)

add_heading_styled('5.4 Utility Modules', 2)
add_table(
    ['Module', 'Responsibility'],
    [
        ['tools.py', 'Shared SQL execution helpers'],
        ['trigger_airflow_dag.py', 'SSH/subprocess Airflow DAG trigger'],
        ['extract_input_template_S3.py', 'Download contract input file from S3'],
        ['readEncryptedConfig.py', 'Decrypt database credentials (Fernet)'],
        ['readMetadata.py', 'Read AWS credentials from database'],
        ['decryption.py', 'Fernet symmetric decryption utility'],
    ]
)

# ==================== 6. DATA FLOW ====================
add_heading_styled('6. Data Flow', 1)

add_heading_styled('6.1 End-to-End Processing Flow', 2)

steps = [
    ('Step 1: Contract Submission (Panel 1)',
     'User selects input type (manual/file), enters contract names, provides delivery name, '
     'confirms details, and triggers the pipeline.'),
    ('Step 2: Metadata Update',
     'Backend truncates admin_fee_metadata table and inserts new contracts with STATUS = 1.'),
    ('Step 3: Pipeline Execution',
     'Backend SSHs to Airflow EC2 instance and triggers DAG: execute_adminFee_Data_Pipeline_v1. '
     'Airflow processes contracts. Status transitions: 1 (Ready) → 2 (SQL Done) → 0 (Completed).'),
    ('Step 4: Status Monitoring (Panel 2)',
     'Frontend polls /api/status/summary every 30 seconds. Displays progress bar and contract table. '
     'User can ask AI questions about status. Fires callback when all contracts reach STATUS = 0.'),
    ('Step 5: Analysis & Reconciliation (Panel 3)',
     'Auto-triggered when processing completes. Normalizes delivery name via LLM, fetches contracts, '
     'runs reconciliation query (PO vs Report), displays summary, and accepts follow-up NL questions.'),
]
for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    doc.add_paragraph(desc)

add_heading_styled('6.2 LLM Interaction Flow', 2)
doc.add_paragraph('Each natural language question goes through a 4-stage LLM pipeline:')
add_table(
    ['Stage', 'Input', 'Output', 'Prompt Template'],
    [
        ['1. SQL Generation', 'User question + schema', 'MySQL SELECT query', 'sql_generator.txt'],
        ['2. SQL Execution', 'Generated SQL', '{columns, rows}', 'N/A (SQLAlchemy)'],
        ['3. Result Analysis', 'Question + raw results', 'Technical analysis', 'contract_analyst_prompt.txt'],
        ['4. Response Formatting', 'Technical analysis', 'Business-friendly answer', 'response_formatter.txt'],
    ]
)

# ==================== 7. DATABASE ====================
add_heading_styled('7. Database Design (High Level)', 1)

add_heading_styled('7.1 Database Instances', 2)
add_table(
    ['Database', 'Purpose'],
    [
        ['joblog_metadata', 'System metadata, audit trails, configuration'],
        ['admin_fee', 'Contract data, PO masters, reports'],
    ]
)

add_heading_styled('7.2 Key Tables', 2)
add_table(
    ['Table', 'Database', 'Purpose'],
    [
        ['admin_fee_metadata', 'joblog_metadata', 'Processing status tracking per contract'],
        ['adminfee_audit_data', 'joblog_metadata', 'Audit trail with file paths and record counts'],
        ['metadata_table_database', 'joblog_metadata', 'AWS credential and config storage'],
        ['PO_Master_{CONTRACT}_{VER}', 'admin_fee', 'Per-contract purchase order spend data'],
        ['admin_fee_report', 'admin_fee', 'Aggregated admin fee report data'],
    ]
)

add_heading_styled('7.3 Status State Machine', 2)
add_table(
    ['Status Code', 'Label', 'Description'],
    [
        ['1', 'Ready', 'Contract inserted, awaiting processing'],
        ['2', 'SQL Done', 'SQL execution complete, Excel pending'],
        ['0', 'Completed', 'Fully processed, Excel generated'],
    ]
)

# ==================== 8. API DESIGN ====================
add_heading_styled('8. API Design (High Level)', 1)

add_heading_styled('8.1 Endpoint Groups', 2)
add_table(
    ['Group', 'Base Path', 'Purpose'],
    [
        ['Agent', '/api/agent/*', 'Processing chat workflow'],
        ['Status', '/api/status/*', 'Live status monitoring'],
        ['Analysis', '/api/analysis/*', 'NL contract analysis'],
        ['Reports', '/api/reports/*', 'Reconciliation & audit reports'],
        ['Health', '/api/health', 'Service health check'],
    ]
)

add_heading_styled('8.2 Endpoint Summary', 2)
add_table(
    ['Method', 'Endpoint', 'Description'],
    [
        ['POST', '/api/agent/start', 'Start new processing session'],
        ['POST', '/api/agent/chat', 'Continue chat conversation'],
        ['GET', '/api/status/summary', 'Get total/completed/in-progress counts'],
        ['GET', '/api/status/contracts', 'Get contract-level status details'],
        ['POST', '/api/status/ask', 'AI-powered status Q&A'],
        ['GET', '/api/analysis/deliveries', 'List available deliveries'],
        ['POST', '/api/analysis/setup', 'Initialize analysis session'],
        ['POST', '/api/analysis/ask', 'NL question answering'],
        ['POST', '/api/reports/reconciliation', 'Generate reconciliation report'],
        ['GET', '/api/reports/audit', 'Get audit trail data'],
        ['GET', '/api/health', 'Health check'],
    ]
)

# ==================== 9. SECURITY ====================
add_heading_styled('9. Security Architecture', 1)

add_heading_styled('9.1 Credential Management', 2)
doc.add_paragraph(
    'Database credentials are stored encrypted using Fernet symmetric encryption. '
    'The decryption flow is: Read key path from Paths.xls → Load Fernet key from disk → '
    'Decrypt credential CSV → Extract host, user, password, database.'
)

add_heading_styled('9.2 Security Controls', 2)
add_table(
    ['Layer', 'Mechanism'],
    [
        ['Encryption', 'Fernet symmetric encryption for database credentials'],
        ['SSH', 'PEM key-based authentication for Airflow trigger'],
        ['DB Connection', 'SQLAlchemy connection pooling with pre-ping validation'],
        ['API', 'CORS middleware with configurable allowed origins'],
        ['Queries', 'Parameterized queries via SQLAlchemy to prevent SQL injection'],
    ]
)

# ==================== 10. DEPLOYMENT ====================
add_heading_styled('10. Deployment Architecture', 1)

add_heading_styled('10.1 Runtime Environment', 2)
add_table(
    ['Service', 'Port', 'Technology'],
    [
        ['Frontend Dev Server', '3000', 'Vite (React)'],
        ['Backend API Server', '8000', 'Uvicorn (FastAPI)'],
        ['Database', '3306', 'AWS RDS MySQL'],
        ['Airflow', '22 (SSH)', 'Apache Airflow on EC2'],
    ]
)

add_heading_styled('10.2 Startup Scripts', 2)
add_table(
    ['Platform', 'Script', 'Process'],
    [
        ['Linux', 'start.sh', 'Runs backend + frontend in background, traps Ctrl+C'],
        ['Windows', 'start.bat', 'Validates deps, installs packages, opens separate CMD windows'],
    ]
)

# ==================== 11. NFR ====================
add_heading_styled('11. Non-Functional Requirements', 1)

add_heading_styled('11.1 Performance', 2)
add_table(
    ['Metric', 'Target'],
    [
        ['API Response Time', '< 2 seconds'],
        ['LLM Response Time', '< 10 seconds'],
        ['Status Polling Interval', '30 seconds'],
        ['DB Connection Pool Recycle', '3600 seconds'],
    ]
)

add_heading_styled('11.2 Scalability', 2)
items = [
    'Session-based state (in-memory) — suitable for single-instance deployment',
    'Database connection pooling for concurrent queries',
    'Stateless API design (except session store)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('11.3 Availability', 2)
items = [
    'Health check endpoint (/api/health) for monitoring',
    'Database pre-ping for connection validation',
    'Graceful error handling with user-friendly messages',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('11.4 Maintainability', 2)
items = [
    'Prompt templates externalized in /prompts directory',
    'Modular backend (database, LLM, API layers separated)',
    'Reusable frontend components with clear prop interfaces',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ==================== SAVE ====================
output_dir = '/home/user/adminFeeAgent2/docs'
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, 'HLD_Document.docx')
doc.save(output_path)
print(f'HLD Word document saved to: {output_path}')
